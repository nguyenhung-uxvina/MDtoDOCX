#!/usr/bin/env python3
"""
Markdown to DOCX converter
Converts Markdown files to Microsoft Word DOCX format
"""

import argparse
import logging
import re
import sys
from pathlib import Path
from html.parser import HTMLParser
from urllib.request import urlopen
from io import BytesIO

import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from PIL import Image as PILImage

# Set up logger
logger = logging.getLogger('md_to_docx')


def setup_logging(verbosity='normal'):
    """
    Configure logging for the application

    Args:
        verbosity: Logging level - 'quiet', 'normal', or 'verbose'
    """
    # Set up log format
    log_format = '%(asctime)s - %(levelname)s - %(message)s'
    date_format = '%Y-%m-%d %H:%M:%S'

    # Determine log level based on verbosity
    if verbosity == 'verbose':
        log_level = logging.DEBUG
    elif verbosity == 'quiet':
        log_level = logging.WARNING
    else:  # normal
        log_level = logging.INFO

    # Configure logging
    logging.basicConfig(
        level=log_level,
        format=log_format,
        datefmt=date_format,
        stream=sys.stderr  # Log to stderr to keep stdout clean
    )

    logger.setLevel(log_level)
    logger.debug(f"Logging initialized at {logging.getLevelName(log_level)} level")


class MarkdownToDocxConverter(HTMLParser):
    """Converts HTML (from Markdown) to DOCX format"""

    def __init__(self, doc):
        super().__init__()
        logger.debug("Initializing HTML to DOCX converter")
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        self.list_level = 0
        self.ordered_list = False
        self.list_counters = [0] * 10  # Support up to 10 nesting levels
        self.in_code_block = False
        self.in_table = False
        self.table = None
        self.current_row = None
        self.current_cell = None
        self.bold = False
        self.italic = False
        self.code = False
        self.heading_level = 0
        self.in_blockquote = False
        self.strikethrough = False
        self.in_task_list = False
        self.superscript = False
        self.subscript = False
        self.highlight = False

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        logger.debug(f"Processing start tag: <{tag}>")

        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.heading_level = int(tag[1])
            self.current_paragraph = self.doc.add_heading('', level=self.heading_level)
            self.current_run = None

        elif tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()
            self.current_run = None

        elif tag == 'strong' or tag == 'b':
            self.bold = True

        elif tag == 'em' or tag == 'i':
            self.italic = True

        elif tag == 'code':
            self.code = True

        elif tag == 'pre':
            self.in_code_block = True
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.style = 'No Spacing'
            self.current_run = None

        elif tag == 'ul':
            self.ordered_list = False
            self.list_level += 1

        elif tag == 'ol':
            self.ordered_list = True
            self.list_level += 1
            self.list_counters[self.list_level - 1] = 0

        elif tag == 'li':
            if self.ordered_list:
                self.list_counters[self.list_level - 1] += 1
                counter = self.list_counters[self.list_level - 1]
                self.current_paragraph = self.doc.add_paragraph(f'{counter}. ', style='List Number')
            else:
                self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
            self.current_run = None

        elif tag == 'a':
            # Links - we'll just add the text and URL
            self.link_url = attrs_dict.get('href', '')

        elif tag == 'img':
            self._handle_image(attrs_dict)

        elif tag == 'table':
            self.in_table = True
            self.table_data = []
            self.current_table_row = []

        elif tag == 'tr':
            self.current_table_row = []

        elif tag in ['th', 'td']:
            self.current_cell_text = []

        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run('\n')

        elif tag == 'blockquote':
            self.in_blockquote = True
            self.current_paragraph = self.doc.add_paragraph()
            # Try to use built-in style, fallback to manual formatting
            try:
                self.current_paragraph.style = 'Intense Quote'
            except KeyError:
                # Fallback: manual blockquote styling
                self.current_paragraph.paragraph_format.left_indent = Inches(0.5)
                self.current_paragraph.paragraph_format.right_indent = Inches(0.5)
                self.current_paragraph.paragraph_format.space_before = Pt(6)
                self.current_paragraph.paragraph_format.space_after = Pt(6)
            self.current_run = None

        elif tag == 'hr':
            # Horizontal rule
            self._add_horizontal_rule()

        elif tag == 'del' or tag == 's':
            # Strikethrough
            self.strikethrough = True

        elif tag == 'sup':
            # Superscript
            self.superscript = True

        elif tag == 'sub':
            # Subscript
            self.subscript = True

        elif tag == 'mark':
            # Highlighting
            self.highlight = True

        elif tag == 'div':
            # Check for page break
            if attrs_dict.get('class') == 'pagebreak':
                self._add_page_break()

    def handle_endtag(self, tag):
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.heading_level = 0
            self.current_paragraph = None
            self.current_run = None

        elif tag == 'p':
            self.current_paragraph = None
            self.current_run = None

        elif tag == 'strong' or tag == 'b':
            self.bold = False
            self.current_run = None

        elif tag == 'em' or tag == 'i':
            self.italic = False
            self.current_run = None

        elif tag == 'code':
            self.code = False
            self.current_run = None

        elif tag == 'pre':
            self.in_code_block = False
            self.current_paragraph = None
            self.current_run = None

        elif tag == 'ul' or tag == 'ol':
            self.list_level -= 1
            if tag == 'ol':
                self.list_counters[self.list_level] = 0

        elif tag == 'li':
            self.current_paragraph = None
            self.current_run = None

        elif tag == 'a':
            if hasattr(self, 'link_url') and self.link_url:
                if self.current_paragraph and self.current_run:
                    # Add URL in parentheses
                    url_run = self.current_paragraph.add_run(f' ({self.link_url})')
                    url_run.font.color.rgb = RGBColor(0, 0, 255)
                    url_run.font.size = Pt(9)
                self.link_url = ''
            self.current_run = None

        elif tag in ['th', 'td']:
            self.current_table_row.append(''.join(self.current_cell_text))

        elif tag == 'tr':
            if self.current_table_row:
                self.table_data.append(self.current_table_row)
            self.current_table_row = []

        elif tag == 'table':
            self._create_table()
            self.in_table = False

        elif tag == 'blockquote':
            self.in_blockquote = False
            self.current_paragraph = None
            self.current_run = None

        elif tag == 'del' or tag == 's':
            self.strikethrough = False
            self.current_run = None

        elif tag == 'sup':
            self.superscript = False
            self.current_run = None

        elif tag == 'sub':
            self.subscript = False
            self.current_run = None

        elif tag == 'mark':
            self.highlight = False
            self.current_run = None

    def handle_data(self, data):
        if not data.strip() and not self.in_code_block:
            return

        if self.in_table and hasattr(self, 'current_cell_text'):
            self.current_cell_text.append(data)
            return

        if not self.current_paragraph:
            self.current_paragraph = self.doc.add_paragraph()

        # Check if we need a new run due to formatting changes
        formatting_changed = (
            not self.current_run or
            self.bold != getattr(self.current_run, '_bold', False) or
            self.italic != getattr(self.current_run, '_italic', False) or
            self.strikethrough != getattr(self.current_run, '_strikethrough', False) or
            self.superscript != getattr(self.current_run, '_superscript', False) or
            self.subscript != getattr(self.current_run, '_subscript', False) or
            self.highlight != getattr(self.current_run, '_highlight', False)
        )

        if formatting_changed:
            self.current_run = self.current_paragraph.add_run(data)
            self.current_run._bold = self.bold
            self.current_run._italic = self.italic
            self.current_run._strikethrough = self.strikethrough
            self.current_run._superscript = self.superscript
            self.current_run._subscript = self.subscript
            self.current_run._highlight = self.highlight

            # Apply formatting
            if self.bold:
                self.current_run.bold = True
            if self.italic:
                self.current_run.italic = True
            if self.strikethrough:
                self.current_run.font.strike = True
            if self.superscript:
                self.current_run.font.superscript = True
            if self.subscript:
                self.current_run.font.subscript = True
            if self.highlight:
                self.current_run.font.highlight_color = 7  # Yellow highlight
            if self.code or self.in_code_block:
                self.current_run.font.name = 'Courier New'
                self.current_run.font.size = Pt(10)
                if not self.in_code_block:
                    self.current_run.font.color.rgb = RGBColor(200, 0, 0)
        else:
            self.current_run.text += data

    def _handle_image(self, attrs):
        """Handle image insertion"""
        src = attrs.get('src', '')
        if not src:
            logger.warning("Image tag found with no src attribute")
            return

        logger.info(f"Processing image: {src}")

        try:
            # Add a new paragraph for the image
            if self.current_paragraph and self.current_paragraph.text:
                img_paragraph = self.doc.add_paragraph()
            else:
                img_paragraph = self.current_paragraph or self.doc.add_paragraph()

            img_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Load image
            if src.startswith(('http://', 'https://')):
                # Download image from URL
                logger.debug(f"Downloading image from URL: {src}")
                response = urlopen(src)
                image_data = BytesIO(response.read())
                logger.debug("Image downloaded successfully")
            else:
                # Local file
                logger.debug(f"Loading local image: {src}")
                image_data = src

            # Add image to document with reasonable size
            run = img_paragraph.add_run()
            picture = run.add_picture(image_data, width=Inches(4.0))
            logger.debug("Image added to document")

            # Add alt text if available
            alt_text = attrs.get('alt', '')
            if alt_text:
                caption_para = self.doc.add_paragraph(alt_text)
                caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption_para.runs[0].italic = True
                caption_para.runs[0].font.size = Pt(9)
                logger.debug(f"Added image caption: {alt_text}")

            self.current_paragraph = None
            self.current_run = None

        except Exception as e:
            # If image loading fails, just add the alt text or filename
            logger.warning(f"Failed to load image '{src}': {e}")
            fallback_text = attrs.get('alt', f'[Image: {src}]')
            if not self.current_paragraph:
                self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.add_run(f'[{fallback_text}]').italic = True
            logger.debug(f"Added fallback text for image: {fallback_text}")

    def _create_table(self):
        """Create a table from collected data"""
        if not self.table_data:
            logger.debug("No table data to create")
            return

        rows = len(self.table_data)
        cols = max(len(row) for row in self.table_data) if self.table_data else 0

        if rows == 0 or cols == 0:
            logger.warning("Table has no rows or columns")
            return

        logger.info(f"Creating table with {rows} rows and {cols} columns")

        table = self.doc.add_table(rows=rows, cols=cols)
        # Try to use built-in style, fallback to safe default
        try:
            table.style = 'Light Grid Accent 1'
            logger.debug("Applied 'Light Grid Accent 1' table style")
        except KeyError:
            try:
                table.style = 'Table Grid'  # Safe fallback
                logger.debug("Applied 'Table Grid' table style (fallback)")
            except KeyError:
                logger.debug("Using default table style")
                pass  # Use default table style

        for i, row_data in enumerate(self.table_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text.strip()
                # Make header row bold
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        logger.debug("Table created successfully")
        self.table_data = []

    def _add_horizontal_rule(self):
        """Add a horizontal rule (line) to the document"""
        # Add a paragraph with a bottom border to simulate a horizontal line
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        # Add a run with a border
        run = para.add_run()
        run.add_text('_' * 80)
        run.font.color.rgb = RGBColor(192, 192, 192)
        run.font.size = Pt(1)

        self.current_paragraph = None
        self.current_run = None

    def _add_page_break(self):
        """Add a page break to the document"""
        # Add a paragraph and insert a page break
        if self.current_paragraph:
            self.current_paragraph.add_run().add_break(WD_BREAK.PAGE)
        else:
            para = self.doc.add_paragraph()
            para.add_run().add_break(WD_BREAK.PAGE)

        self.current_paragraph = None
        self.current_run = None


def _preprocess_task_lists(content):
    """Convert task list syntax to a format we can parse"""
    # Convert [ ] to unchecked checkbox symbol
    content = re.sub(r'^(\s*)[-*]\s+\[ \]\s+', r'\1- ☐ ', content, flags=re.MULTILINE)
    # Convert [x] or [X] to checked checkbox symbol
    content = re.sub(r'^(\s*)[-*]\s+\[[xX]\]\s+', r'\1- ☑ ', content, flags=re.MULTILINE)
    return content


def _preprocess_page_breaks(content):
    """Convert page break markers to HTML"""
    # Support multiple page break syntaxes
    content = re.sub(r'<!--\s*pagebreak\s*-->', '<div class="pagebreak"></div>', content, flags=re.IGNORECASE)
    content = re.sub(r'\\pagebreak', '<div class="pagebreak"></div>', content)
    content = re.sub(r'<pagebreak\s*/>', '<div class="pagebreak"></div>', content, flags=re.IGNORECASE)
    return content


def _preprocess_special_formatting(content):
    """Convert special formatting that markdown doesn't handle natively"""
    # Strikethrough: ~~text~~ (convert to HTML del tag) - MUST BE BEFORE SUBSCRIPT
    content = re.sub(r'~~([^~]+)~~', r'<del>\1</del>', content)
    # Highlighting: ==text==
    content = re.sub(r'==([^=]+)==', r'<mark>\1</mark>', content)
    # Superscript: text^superscript^
    content = re.sub(r'\^([^\^]+)\^', r'<sup>\1</sup>', content)
    # Subscript: text~subscript~ (single tilde, after strikethrough is processed)
    content = re.sub(r'(?<!~)~([^~]+)~(?!~)', r'<sub>\1</sub>', content)
    return content


def convert_markdown_to_docx(markdown_file, output_file=None):
    """
    Convert a Markdown file to DOCX format

    Args:
        markdown_file: Path to input Markdown file
        output_file: Path to output DOCX file (optional)
    """
    markdown_path = Path(markdown_file)
    logger.info(f"Starting conversion: {markdown_path}")

    if not markdown_path.exists():
        logger.error(f"File not found: {markdown_file}")
        raise FileNotFoundError(f"Markdown file not found: {markdown_file}")

    # Read markdown content with robust encoding handling
    logger.debug("Reading markdown file")
    try:
        # Try UTF-8 with BOM first (common on Windows)
        with open(markdown_path, 'r', encoding='utf-8-sig') as f:
            md_content = f.read()
        logger.debug("File read successfully with UTF-8-sig encoding")
    except UnicodeDecodeError:
        logger.debug("UTF-8 decoding failed, trying cp1252")
        try:
            # Fallback to Windows default encoding
            with open(markdown_path, 'r', encoding='cp1252') as f:
                md_content = f.read()
            logger.warning("File read with cp1252 encoding (fallback)")
        except UnicodeDecodeError:
            logger.debug("cp1252 decoding failed, trying latin-1")
            # Last resort: latin-1 (never fails but may have wrong chars)
            with open(markdown_path, 'r', encoding='latin-1') as f:
                md_content = f.read()
            logger.warning("File read with latin-1 encoding (last resort)")

    # Convert markdown to HTML
    logger.debug("Converting Markdown to HTML")
    md = markdown.Markdown(extensions=[
        'extra',           # Includes tables, footnotes, attr_list, def_list, fenced_code, abbr
        'codehilite',      # Code highlighting
        'tables',          # Tables
        'fenced_code',     # Fenced code blocks
        'nl2br',           # Newline to <br>
        'sane_lists',      # Better list handling
        'smarty'           # Smart quotes and dashes
    ])

    # Pre-process for special features (not natively supported)
    logger.debug("Preprocessing special formatting")
    md_content = _preprocess_special_formatting(md_content)
    md_content = _preprocess_task_lists(md_content)
    md_content = _preprocess_page_breaks(md_content)

    html_content = md.convert(md_content)
    logger.debug(f"HTML conversion complete ({len(html_content)} characters)")

    # Create DOCX document
    logger.debug("Creating DOCX document")
    doc = Document()

    # Set up document margins and styles
    logger.debug("Setting document margins")
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Parse HTML and convert to DOCX
    logger.info("Parsing HTML and building DOCX structure")
    converter = MarkdownToDocxConverter(doc)
    converter.feed(html_content)

    # Determine output filename
    if output_file is None:
        output_file = markdown_path.with_suffix('.docx')
    else:
        output_file = Path(output_file)

    # Save document
    logger.info(f"Saving document to: {output_file}")
    doc.save(output_file)
    logger.info("Conversion completed successfully")

    return output_file


def main():
    """Main CLI entry point"""
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to DOCX format',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python md_to_docx.py input.md
  python md_to_docx.py input.md -o output.docx
  python md_to_docx.py file1.md file2.md file3.md
  python md_to_docx.py *.md
        """
    )

    parser.add_argument(
        'input',
        nargs='+',
        help='Input Markdown file(s)'
    )

    parser.add_argument(
        '-o', '--output',
        help='Output DOCX file (only for single input file)',
        default=None
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose logging (DEBUG level)'
    )

    parser.add_argument(
        '-q', '--quiet',
        action='store_true',
        help='Quiet mode - only show warnings and errors'
    )

    args = parser.parse_args()

    # Set up logging based on verbosity flags
    if args.verbose:
        setup_logging('verbose')
    elif args.quiet:
        setup_logging('quiet')
    else:
        setup_logging('normal')

    # Validate: -o only works with single file
    if args.output and len(args.input) > 1:
        logger.error("-o/--output can only be used with a single input file")
        print("Error: -o/--output can only be used with a single input file", file=sys.stderr)
        return 1

    # Check for conflicting verbosity flags
    if args.verbose and args.quiet:
        logger.error("Cannot use both --verbose and --quiet flags")
        print("Error: Cannot use both --verbose and --quiet flags", file=sys.stderr)
        return 1

    # Track results for batch operations
    success_count = 0
    failed_files = []

    logger.info(f"Processing {len(args.input)} file(s)")

    try:
        # Process each file
        for input_file in args.input:
            try:
                # Show progress indicator
                input_name = Path(input_file).name
                print(f"Converting '{input_name}'... ", end='', flush=True)

                output_path = convert_markdown_to_docx(input_file, args.output)

                print("[OK]")
                if len(args.input) == 1:
                    # Only show output path for single file
                    print(f"Output: {output_path}")
                success_count += 1

            except FileNotFoundError as e:
                logger.error(f"File not found: {input_file}")
                print("[FAILED]", file=sys.stderr)
                print(f"  Error: {e}", file=sys.stderr)
                failed_files.append((input_file, str(e)))
            except PermissionError as e:
                logger.error(f"Permission denied for {input_file}: {e}")
                print("[FAILED]", file=sys.stderr)
                print(f"  Error: Permission denied - unable to write output file", file=sys.stderr)
                print(f"  Check if the file is open in another program", file=sys.stderr)
                failed_files.append((input_file, "Permission denied"))
            except UnicodeDecodeError:
                logger.error(f"Encoding not supported for {input_file}")
                print("[FAILED]", file=sys.stderr)
                print(f"  Error: File encoding not supported", file=sys.stderr)
                print(f"  Tip: Save as UTF-8 in your text editor", file=sys.stderr)
                failed_files.append((input_file, "Encoding not supported"))
            except Exception as e:
                logger.error(f"Unexpected error processing {input_file}: {e}", exc_info=True)
                print("[FAILED]", file=sys.stderr)
                print(f"  Error: {e}", file=sys.stderr)
                failed_files.append((input_file, str(e)))

        # Show summary for batch operations
        if len(args.input) > 1:
            logger.info(f"Batch conversion complete: {success_count}/{len(args.input)} successful")
            print("\n" + "=" * 50)
            print(f"Conversion Summary:")
            print(f"  Successful: {success_count}/{len(args.input)}")
            if failed_files:
                print(f"  Failed: {len(failed_files)}")
                print("\nFailed files:")
                for filename, error in failed_files:
                    print(f"  - {filename}: {error}")
            print("=" * 50)

        # Return appropriate exit code
        if success_count == 0:
            logger.error("All conversions failed")
            return 1  # All failed
        elif failed_files:
            logger.warning(f"Some conversions failed: {len(failed_files)} out of {len(args.input)}")
            return 2  # Some failed
        else:
            logger.info("All conversions completed successfully")
            return 0  # All succeeded

    except KeyboardInterrupt:
        logger.warning("Conversion cancelled by user")
        print("\n[CANCELLED]", file=sys.stderr)
        if len(args.input) > 1 and success_count > 0:
            print(f"Converted {success_count} file(s) before cancellation", file=sys.stderr)
        return 130


if __name__ == '__main__':
    sys.exit(main())
